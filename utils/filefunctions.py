import os,json,requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH #导入对齐
from docx.shared import Pt,Inches,Cm #导入单位
from docx.oxml.ns import qn #导入中文字体
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT,WD_TABLE_ALIGNMENT,WD_ALIGN_VERTICAL
from utils.config import FILESAVEPATH
from utils.governmentaffairs import MaterialList,MaterialExampleList


def get_path():
    file_path = FILESAVEPATH
    file_name_list = os.listdir(file_path)
    #print(file_name_list)
    return file_name_list


def file_save_name(file_model,governmentaffair):
    file_save_path = is_file_exsit(governmentaffair)
    file_model.save(file_save_path)


def is_file_exsit(governmentaffair):
    result_path = FILESAVEPATH
    folder_path = os.path.join(result_path, f'{governmentaffair.dept_name}')
    if not os.path.exists(folder_path):
        os.mkdir(folder_path)
    service_name = governmentaffair.service_name.replace('/','、').replace('<','').replace('>','')
    save_name = r"{}-{}服务指南.docx".format(governmentaffair.dept_name,service_name)
    file_save_path = os.path.join(folder_path,save_name)
    return file_save_path





def add_title(file_model,governmentaffair):
    title_name = "{}服务指南".format(governmentaffair.service_name)
    #添加段落，设置段落对齐格式
    #paragraph = file_model.add_paragraph()
    paragraph = file_model.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #给段落添加块，文字内容
    title = paragraph.add_run(title_name)
    #设置字体
    title.font.name = "Arial"
    title.element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    title.font.size = Pt(15)
    # #设置加粗
    # title.font.bole = True
    #设置前后段落间距离
    paragraph.paragraph_format.space_after = Pt(15)
    paragraph.paragraph_format.space_before = Pt(15)


def add_title_style(text,file_model):
    paragraph=file_model.add_paragraph()
    content = paragraph.add_run(text)
    content.font.name = "Arial"
    content.element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    content.font.size = Pt(10.5)



def add_content_style(text,file_model):
    paragraph=file_model.add_paragraph()
    content = paragraph.add_run(text)
    content.font.name = '宋体'
    content.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    content.font.size = Pt(10.5)
    paragraph.paragraph_format.first_line_indent = Inches(0.4)


def add_service_codenum(file_model,governmentaffair):
    title = "一、事项编码"
    content = governmentaffair.service_codenum
    add_title_style(title,file_model)
    add_content_style(content,file_model)
    
    
def add_apply_condition(file_model,governmentaffair):
    title = "二、适用范围"
    content = "符合申请{}条件的{}".format(governmentaffair.service_name,governmentaffair.for_user)
    add_title_style(title,file_model)
    add_content_style(content,file_model)


def add_service_type(file_model,governmentaffair):
    title = "三、事项类型"
    content = governmentaffair.service_type
    add_title_style(title,file_model)
    add_content_style(content,file_model)
    

def add_legal_foundation(file_model,governmentaffair):
    title = "四、设立依据"
    add_title_style(title,file_model)
    legal_foundation = governmentaffair.legal_foundation
    legal_foundation_list = legal_foundation.split("\n")
    for content in legal_foundation_list:
        add_content_style(content,file_model)


def add_dept_name(file_model,governmentaffair):
    title1 = "五、受理机构"
    title2 = "六、决定机构"
    content = governmentaffair.dept_name
    add_title_style(title1,file_model)
    add_content_style(content,file_model)
    add_title_style(title2,file_model)
    add_content_style(content,file_model)

    
def add_apply_condition_desc(file_model,governmentaffair):
    title = "七、办理的条件"
    add_title_style(title,file_model)
    apply_condition_desc = governmentaffair.apply_condition_desc
    apply_condition_list = apply_condition_desc.split('；')
    for apply_condition in apply_condition_list:
        if apply_condition == apply_condition_list[-1]:
            content = apply_condition
        else:
            content = ''.join([apply_condition,"；"])
        add_content_style(content,file_model)


#完善盛情材料列表内容
def add_material_list(file_model,governmentaffair):
    title = "八、申请材料"
    add_title_style(title,file_model)
    if governmentaffair.materia_lists == '[]':
        add_content_style('无需提交申请材料',file_model)
    else:
        text = '    {}申请材料'.format(governmentaffair.service_name)
        paragraph=file_model.add_paragraph()
        content = paragraph.add_run(text)
        content.font.name = '宋体'
        content.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
        content.font.size = Pt(10.5)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #增加申请材料表格
        table = file_model.add_table(rows=1,cols=0,style ='Light Grid')
        table.add_column(width=Cm(1))
        table.add_column(width=Cm(4.5))
        table.add_column(width=Cm(3))
        table.add_column(width=Cm(2.5))
        table.add_column(width=Cm(4))
        heading_cells = table.rows[0].cells
        heading_cells[0].text = '序号'
        heading_cells[1].text = '提交材料名称'
        heading_cells[2].text = '材料类型'
        heading_cells[3].text = '来源渠道'
        heading_cells[4].text = '填报须知'
        add_material(table,governmentaffair)    
        table_style(table)


def add_material(table,governmentaffair):
    material_lists = handle_material_lists(governmentaffair)
    for m in material_lists:
        row = table.add_row()
        row.cells[0].text = str(m[0])
        row.cells[1].text = str(m[1])
        row.cells[2].text = str(m[2])
        row.cells[3].text = str(m[3])
        row.cells[4].text = str(m[4])


def table_style(table):
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
                for text in paragraph.runs:
                    text.font.name = '宋体'
                    text.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
                    text.font.size = Pt(10.5)
    table.autofit=True
    table.alignment=WD_TABLE_ALIGNMENT.CENTER

                    
def handle_material_lists(governmentaffair):
    if governmentaffair.materia_lists != '[]':
        material_excel = governmentaffair.materia_lists.replace('\r', '').replace('\n', '').replace('\"','')
        material_content = json.loads(material_excel.replace("'", "\""))
        if len(material_content) > 0:
            material_lists = []
            i=1
            for material in material_content:
                material_num = i
                material_name = material.get('name')
                material_copynum = material.get('copyNum')
                material_papernum = material.get('submitMedium')
                material_request = '原件：{}，复印件：{}；'.format(material_papernum,material_copynum)
                material_source = material.get('srcWay')
                material_accept_require = material.get('acceptRequire')
                material_list = [material_num,material_name,material_request,material_source,material_accept_require]
                material_lists.append(material_list)
                i+=1
    return material_lists

    
def add_accept_method(file_model,governmentaffair):
    title = "九、受理方式"
    add_title_style(title,file_model)
    content_part_one= "（一）窗口受理：直接到县政服务大厅{}窗口提交申请材料。".format(governmentaffair.accept_window)
    add_content_style(content_part_one,file_model)
    content_part_two = "（二）网上申报：进入河南政务服务网（http://was.hnzwfw.gov.cn/ycslypt_web/serviceApply.action）按照提示进行网上申报。"
    add_content_style(content_part_two,file_model)


def add_flow_desc(file_model,governmentaffair):
    title = "十、办理流程"
    add_title_style(title,file_model)
    if governmentaffair.flow_type == "行政许可类":
        add_flow_type_xzxk(file_model,governmentaffair)
    elif governmentaffair.flow_type == "其他类":
        add_flow_tyoe_qt(file_model,governmentaffair)
    elif governmentaffair.flow_type == "即办件":
        add_flow_tyoe_jb(file_model,governmentaffair)


def add_flow_type_xzxk(file_model,governmentaffair):
    content_part_one = '（一）申请'
    content_part_two = '申请人通过政务服务网、移动端和县政务服务大厅{}窗口进行事项的申请，提交有关申请材料和反映真实情况，并对其申请材料实质内容的真实性负责。'.format(governmentaffair.accept_window)
    content_part_three = '（二）受理'
    content_part_four = '1.后台审批区收到申请人提交的材料后，应立即进行审核，符合受理条件的，应予以受理，出具《受理通知书》，并以短信形式通知申请人。'
    content_part_five = '2.申请材料不齐全或者不符合法定形式的，通过{}窗口出具《申请材料补正通知书》，一次性告知申请人需要补正的全部内容和补正期限。'.format(governmentaffair.accept_window)
    content_part_six = '3.申请材料不符合受理条件的，应当做出不予受理决定，并通过{}窗口向申请人出具《不予受理决定书》，列明不予受理的理由。'.format(governmentaffair.accept_window)
    content_part_seven = '（三）审核'
    content_part_eight = '后台审批区对受理的申请材料要进行实质性审核，需要现场踏勘的，应安排相关人员进行现场踏勘，并出具踏勘意见；需要征询相关部门意见的，应征询相关部门意见。通过以上审核提出审批建议。'
    content_part_nine = '（四）决定'
    content_part_ten = '首席代表收到审核人呈报的审批建议后作出准予行政许可或不予行政许可决定，承办人根据首席代表的决定制作《准予行政许可决定书》或《不予行政许可决定书》。'
    content_part_eleven = '（五）送达'
    content_part_twelve = '根据申请人意愿，由{}出证窗口电话通知申请人领取《准予行政许可决定书》或《不予行政许可决定书》；或通过快递邮寄给申请人。'.format(governmentaffair.accept_window)
    add_content_style(content_part_one,file_model)
    add_content_style(content_part_two,file_model)
    add_content_style(content_part_three,file_model)
    add_content_style(content_part_four,file_model)
    add_content_style(content_part_five,file_model)
    add_content_style(content_part_six,file_model)
    add_content_style(content_part_seven,file_model)
    add_content_style(content_part_eight,file_model)
    add_content_style(content_part_nine,file_model)
    add_content_style(content_part_ten,file_model)
    add_content_style(content_part_eleven,file_model)
    add_content_style(content_part_twelve,file_model)


def add_flow_tyoe_qt(file_model,governmentaffair):
    content_part_one = '（一）申请'
    content_part_two = '申请人通过政务服务网、移动端和县政务服务大厅{}窗口进行事项的申请，提交有关申请材料和反映真实情况，并对其申请材料实质内容的真实性负责。'.format(governmentaffair.accept_window)
    content_part_three = '（二）受理'
    content_part_four = '1.后台审批区收到申请人提交的材料后，应立即进行审核，符合受理条件的，应予以受理，出具《受理通知书》，并以短信形式通知申请人。'
    content_part_five = '2.申请材料不齐全或者不符合法定形式的，通过{}窗口出具《申请材料补正通知书》，一次性告知申请人需要补正的全部内容和补正期限。'.format(governmentaffair.accept_window)
    content_part_six = '3.申请材料不符合受理条件的，应当做出不予受理决定，并通过{}窗口向申请人出具《不予受理决定书》，列明不予受理的理由。'.format(governmentaffair.accept_window)
    content_part_seven = '（三）审核'
    content_part_eight = '后台审批区对受理的申请材料要进行实质性审核，需要现场踏勘的，应安排相关人员进行现场踏勘，并出具踏勘意见；需要征询相关部门意见的，应征询相关部门意见。通过以上审核提出审批建议。'
    content_part_nine = '（四）决定'
    content_part_ten = '申请符合规定的，准予审批通过；申请不符合规定的，不予审批通过。'
    content_part_eleven = '（五）送达'
    content_part_twelve = '根据申请人意愿，由{}出证窗口电话通知申请人领取或通过快递邮寄给申请人。'.format(governmentaffair.accept_window)
    add_content_style(content_part_one,file_model)
    add_content_style(content_part_two,file_model)
    add_content_style(content_part_three,file_model)
    add_content_style(content_part_four,file_model)
    add_content_style(content_part_five,file_model)
    add_content_style(content_part_six,file_model)
    add_content_style(content_part_seven,file_model)
    add_content_style(content_part_eight,file_model)
    add_content_style(content_part_nine,file_model)
    add_content_style(content_part_ten,file_model)
    add_content_style(content_part_eleven,file_model)
    add_content_style(content_part_twelve,file_model)


def add_flow_tyoe_jb(file_model,governmentaffair):
    content_part_one = '（一）申请'
    content_part_two = '申请人通过政务服务网、移动端和县政务服务大厅{}窗口进行事项的申请，提交有关申请材料和反映真实情况，并对其申请材料实质内容的真实性负责。'.format(governmentaffair.accept_window)
    content_part_three = '（二）受理'
    content_part_four = '1.后台审批区收到申请人提交的材料后，应立即进行审核，符合受理条件的，应予以受理，出具《受理通知书》，并以短信形式通知申请人。'
    content_part_five = '2.申请材料不齐全或者不符合法定形式的，通过{}窗口出具《申请材料补正通知书》，一次性告知申请人需要补正的全部内容和补正期限。'.format(governmentaffair.accept_window)
    content_part_six = '3.申请材料不符合受理条件的，应当做出不予受理决定，并通过{}窗口向申请人出具《不予受理决定书》，列明不予受理的理由。'.format(governmentaffair.accept_window)
    content_part_seven = '（三）决定'
    content_part_eight = '申请符合规定的，准予审批通过；申请不符合规定的，不予审批通过。申请结果当场告知申请人。'
    add_content_style(content_part_one,file_model)
    add_content_style(content_part_two,file_model)
    add_content_style(content_part_three,file_model)
    add_content_style(content_part_four,file_model)
    add_content_style(content_part_five,file_model)
    add_content_style(content_part_six,file_model)
    add_content_style(content_part_seven,file_model)
    add_content_style(content_part_eight,file_model)


def add_time_limit(file_model,governmentaffair):
    title = "十一、办理时限"
    add_title_style(title,file_model)
    content_part_one = "（一）法定时限"
    add_content_style(content_part_one,file_model)
    content_part_two = "{}个工作日。".format(governmentaffair.legal_days)
    add_content_style(content_part_two,file_model)
    content_part_three = "（二）承诺时限"
    add_content_style(content_part_three,file_model)
    content_part_four = "{}个工作日。".format(governmentaffair.promise_days)
    add_content_style(content_part_four,file_model)


def add_charging_basis(file_model,governmentaffair):
    title = '十二、收费依据及标准'
    add_title_style(title,file_model)
    content = '不收费。'
    add_content_style(content,file_model)


def add_result_get_method(file_model,governmentaffair):
    title = '十三、结果送达'
    add_title_style(title,file_model)
    content = '根据申请人意愿直接送达或邮寄送达。'
    add_content_style(content,file_model)


def add_administrative_relief(file_model,governmentaffair):
    title = '十四、行政救济途径与方式'
    add_title_style(title,file_model)
    content_part_one = '（一）申请人在办理政务服务事项的过程中，依法享有陈述权、申辩权；'
    content_part_two = '（二）申请人的申请被驳回的有权要求说明理由；'
    content_part_three = '（三）申请人不服办理结果的，有权在收到行政许可决定之日起60日内向修武县人民政府申请行政复议，或者在6个月内向修武县人民法院提起行政诉讼。'
    add_content_style(content_part_one,file_model)
    add_content_style(content_part_two,file_model)
    add_content_style(content_part_three,file_model)


def add_consult_method(file_model,governmentaffair):
    title = '十五、咨询方式'
    add_title_style(title,file_model)
    content_part_one = '（一）现场咨询'
    content_part_two = '修武县政务服务大厅一楼咨询投诉台。'
    content_part_three = '（二）电话咨询'
    content_part_four = '0391-7180080  {}'.format(governmentaffair.consult_tel)
    content_part_five = '（三）网上咨询'
    content_part_six = '河南政务服务网，网址为https://www.hnzwfw.gov.cn'
    add_content_style(content_part_one,file_model)
    add_content_style(content_part_two,file_model)
    add_content_style(content_part_three,file_model)
    add_content_style(content_part_four,file_model)
    add_content_style(content_part_five,file_model)
    add_content_style(content_part_six,file_model)


def add_complaint_method(file_model,governmentaffair):
    title = '十六、监督投诉渠道'
    add_title_style(title,file_model)
    content_part_one = '（一）现场监督投诉'
    content_part_two = '修武县政务服务大厅一楼咨询投诉台。'
    content_part_three = '（二）电话监督投诉'
    content_part_four = '1.单位：{}'.format(governmentaffair.complaint_tel)
    content_part_five = '2.修武县政务服务和大数据管理局投诉电话：0391-7180080'
    content_part_six = '（三）网上监督投诉'
    content_part_seven = '河南政务服务网，网址为 https://www.hnzwfw.gov.cn'
    add_content_style(content_part_one,file_model)
    add_content_style(content_part_two,file_model)
    add_content_style(content_part_three,file_model)
    add_content_style(content_part_four,file_model)
    add_content_style(content_part_five,file_model)
    add_content_style(content_part_six,file_model)
    add_content_style(content_part_seven,file_model)


def add_processing_address(file_model,governmentaffair):
    title = '十七、办理地址和时间'
    add_title_style(title,file_model)
    content_part_one = '地址：修武县宁城路69号县政务服务中心{}受理窗口'.format(governmentaffair.accept_window)
    content_part_two = '时间：每周一至周五（法定节假日除外）'
    content_part_three = '     夏季 上午8:00—12:00 下午15:00—18:00'
    content_part_four = '     冬季 上午8:00—12:00 下午14:30—17:30'
    add_content_style(content_part_one,file_model)
    add_content_style(content_part_two,file_model)
    add_content_style(content_part_three,file_model)
    add_content_style(content_part_four,file_model)


def add_process_results(file_model,governmentaffair):
    title = '十八、办理进程和结果查询'
    add_title_style(title,file_model)
    content_part_one = '1.现场查询'
    content_part_two = '修武县政务服务中心一楼自助服务区自助查询机或{}受理窗口。'.format(governmentaffair.accept_window)
    content_part_three = '2.电话查询'
    content_part_four = governmentaffair.consult_tel
    content_part_five = '3.网上查询'
    content_part_six = '河南政务服务网，网址为https://www.hnzwfw.gov.cn'
    add_content_style(content_part_one,file_model)
    add_content_style(content_part_two,file_model)
    add_content_style(content_part_three,file_model)
    add_content_style(content_part_four,file_model)
    add_content_style(content_part_five,file_model)
    add_content_style(content_part_six,file_model)


def add_result(file_model,governmentaffair):
    title = '十九、办理结果样本'
    add_title_style(title,file_model)
    content = governmentaffair.result_name
    add_content_style(content,file_model)


def add_enclosure(file_model,governmentaffair):
    title = '二十、附件'
    add_title_style(title,file_model)
    content_part_one = '附件1：事项流程图'
    content_part_two = '附件2：{}'.format(governmentaffair.result_name)
    content_part_three = '附件3：申请材料'
    add_content_style(content_part_one,file_model)
    add_content_style(content_part_two,file_model)
    add_content_style(content_part_three,file_model)


def modify_pagesettings(file_model):
    section = file_model.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    for paragraph in file_model.paragraphs[1:]:
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.space_before = Pt(6)


def down_material_lists(governmentaffair):
    if governmentaffair.materia_lists != '[]':
        material_excel = governmentaffair.materia_lists.replace('\r', '').replace('\n', '').replace('\"','')
        material_content = json.loads(material_excel.replace("'", "\""))
        if len(material_content) > 0:
            i=1
            for material in material_content:
                material_num = i
                material_service_name = governmentaffair.service_name.replace('/','、').replace('<','').replace('>','')
                material_service_unid = governmentaffair.service_unid
                material_dept_name = governmentaffair.dept_name
                material_name = material.get('name').replace('/','、')
                if material_name:
                    material_name = check_file_name(material_name)
                material_file_url = material.get('fileUrl')
                if material_file_url:
                    material_list = MaterialList(material_num,material_service_name,material_service_unid,material_dept_name,material_name,material_file_url)
                    material_list.material_save_path()
                    request_material_file_down(material_list.material_file_url,material_list.service_name_path,material_list.material_save_file)                
                #下载申请材料实例文本
                # material_example_file_url = material.get('exampleFileUrl')
                # if material_example_file_url:
                #     material_example_list = MaterialExampleList(material_num,material_service_name,material_service_unid,material_dept_name,material_name,material_example_file_url)
                #     material_example_list.material_save_path()
                #     request_material_file_down(material_example_list.material_example_file_url,material_example_list.service_name_path,material_example_list.material_example_save_file)
                i+=1


def request_material_file_down(file_ulr,file_path,file_name):
    file = os.path.exists(file_name)
    if not file:
        headers = {
                'User-Agent':'Mozilla/5.0 (Windows NT 6.1; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.77 Safari/537.36 Edg/91.0.864.37',
                }
        resp = requests.get(file_ulr,headers = headers,stream=True)
        temporary_file = os.path.join(file_path,'temporary.text')
        with open(temporary_file, 'wb') as f:
            for chunk in resp.iter_content(chunk_size=1024 * 1024):
                if chunk:
                    f.write(chunk)        
        os.rename(temporary_file, file_name)
        print("{}保存完毕".format(file_name))


def check_file_name(material_name):
    sets = ['/', '\\', ':', '*', '?', '"', '<', '>', '|']
    for char in material_name:
        if char in sets:
            material_name = material_name.replace(char, '')
    if len(material_name)>100:
        material_name = material_name[:99]
    return material_name


def add_contents(file_model,governmentaffair):
    add_title(file_model,governmentaffair)
    add_service_codenum(file_model,governmentaffair)
    add_apply_condition(file_model,governmentaffair)
    add_service_type(file_model,governmentaffair)
    add_legal_foundation(file_model,governmentaffair)
    add_dept_name(file_model,governmentaffair)
    add_apply_condition_desc(file_model,governmentaffair)
    add_material_list(file_model,governmentaffair)
    add_accept_method(file_model,governmentaffair)
    add_flow_desc(file_model,governmentaffair)
    add_time_limit(file_model,governmentaffair)
    add_charging_basis(file_model,governmentaffair)
    add_result_get_method(file_model,governmentaffair)
    add_administrative_relief(file_model,governmentaffair)
    add_consult_method(file_model,governmentaffair)
    add_complaint_method(file_model,governmentaffair)
    add_processing_address(file_model,governmentaffair)
    add_process_results(file_model,governmentaffair)
    add_result(file_model,governmentaffair)
    add_enclosure(file_model,governmentaffair)    
    modify_pagesettings(file_model)